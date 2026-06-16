"""Parser tests. Pure Python, no API keys needed."""
import io

import pandas as pd
import pytest

from app.parsers.parse import (
    parse_file,
    parse_txt,
    parse_csv,
    parse_xlsx,
    parse_docx,
    parse_pdf,
)


def test_txt():
    assert parse_txt(b"hello world") == "hello world"


def test_txt_bad_bytes_no_crash():
    # invalid utf-8 should not raise
    out = parse_txt(b"\xff\xfe abc")
    assert "abc" in out


def test_csv_rows_and_headers():
    csv = b"name,age\nAlice,30\nBob,25"
    out = parse_csv(csv)
    assert "name: Alice | age: 30" in out
    assert "name: Bob | age: 25" in out


def test_csv_empty():
    assert parse_csv(b"") == ""


def test_xlsx_multi_sheet():
    buf = io.BytesIO()
    with pd.ExcelWriter(buf, engine="openpyxl") as w:
        pd.DataFrame({"x": [1], "y": [2]}).to_excel(w, sheet_name="S1", index=False)
        pd.DataFrame({"a": ["foo"]}).to_excel(w, sheet_name="S2", index=False)
    out = parse_xlsx(buf.getvalue())
    assert "[Sheet: S1]" in out
    assert "[Sheet: S2]" in out
    assert "x: 1 | y: 2" in out
    assert "a: foo" in out


def test_docx_paragraphs_and_tables():
    from docx import Document

    buf = io.BytesIO()
    doc = Document()
    doc.add_paragraph("Executive summary")
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Metric"
    table.cell(0, 1).text = "Value"
    table.cell(1, 0).text = "Growth"
    table.cell(1, 1).text = "12%"
    doc.save(buf)

    out = parse_docx(buf.getvalue())
    assert "Executive summary" in out
    assert "[Table: 1]" in out
    assert "Growth | 12%" in out


def test_pdf_empty_no_crash():
    from pypdf import PdfWriter

    buf = io.BytesIO()
    writer = PdfWriter()
    writer.add_blank_page(width=72, height=72)
    writer.write(buf)

    assert parse_pdf(buf.getvalue()) == ""


def test_dispatch_by_extension():
    assert parse_file("notes.txt", b"hi") == "hi"
    assert parse_file("conversation.md", b"# Export\n\nhello") == "# Export\n\nhello"


def test_unsupported_type_raises():
    with pytest.raises(ValueError):
        parse_file("slides.pptx", b"PK")
