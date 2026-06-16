"""File parsers. txt / md / csv / xlsx / pdf / docx -> single text string.

Uniform output so everything downstream is format-agnostic.
"""
import io

import pandas as pd

SUPPORTED_EXTENSIONS = {".txt", ".md", ".csv", ".xlsx", ".pdf", ".docx"}


def _ext(filename: str) -> str:
    """Lowercase extension with leading dot, or '' if none."""
    if "." not in filename:
        return ""
    return "." + filename.rsplit(".", 1)[-1].lower()


def _rows_to_text(df: pd.DataFrame) -> str:
    """One line per row: 'col: val | col: val'. Preserves headers as labels."""
    lines = []
    for _, row in df.iterrows():
        line = " | ".join(f"{col}: {val}" for col, val in row.items())
        lines.append(line)
    return "\n".join(lines)


def parse_txt(data: bytes) -> str:
    return data.decode("utf-8", errors="replace")


def parse_csv(data: bytes) -> str:
    try:
        df = pd.read_csv(io.BytesIO(data))
    except pd.errors.EmptyDataError:
        return ""
    return _rows_to_text(df)


def parse_xlsx(data: bytes) -> str:
    """All sheets. Each tagged with its sheet name so chunks stay attributable."""
    xl = pd.ExcelFile(io.BytesIO(data))
    parts = []
    for sheet in xl.sheet_names:
        df = xl.parse(sheet)
        if df.empty:
            continue
        parts.append(f"[Sheet: {sheet}]")
        parts.append(_rows_to_text(df))
    return "\n".join(parts)


def parse_pdf(data: bytes) -> str:
    """Extract page text from a PDF, tagging each page for later citation."""
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(data))
    parts = []
    for i, page in enumerate(reader.pages, start=1):
        text = page.extract_text() or ""
        text = text.strip()
        if text:
            parts.append(f"[Page: {i}]\n{text}")
    return "\n\n".join(parts)


def parse_docx(data: bytes) -> str:
    """Extract paragraphs and table rows from a Word document."""
    from docx import Document

    doc = Document(io.BytesIO(data))
    parts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    for table_i, table in enumerate(doc.tables, start=1):
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                rows.append(" | ".join(cells))
        if rows:
            parts.append(f"[Table: {table_i}]\n" + "\n".join(rows))
    return "\n".join(parts)


def parse_file(filename: str, data: bytes) -> str:
    """Dispatch by extension. Raises ValueError on unsupported type."""
    ext = _ext(filename)
    if ext in {".txt", ".md"}:
        return parse_txt(data)
    if ext == ".csv":
        return parse_csv(data)
    if ext == ".xlsx":
        return parse_xlsx(data)
    if ext == ".pdf":
        return parse_pdf(data)
    if ext == ".docx":
        return parse_docx(data)
    raise ValueError(
        f"Unsupported file type '{ext or filename}'. "
        f"Allowed: {sorted(SUPPORTED_EXTENSIONS)}"
    )
